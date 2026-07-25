#!/usr/bin/env python3
"""Publish due Bibel-Worte packages from a private input checkout.

Input convention:
  inhalte/YYMMDD Thema.docx
  inhalte/YYMMDD Thema DEU.jpg
  inhalte/YYMMDD Thema ENG.jpg
  inhalte/YYMMDD Thema FR.jpg

The DOCX contains the fixed DEU/ENG/FR headings used by the Bibel-Worte
multilingual template. The six-digit filename prefix is authoritative for the
release date. Documents dated after the cutoff are ignored. Existing published
items remain in the manifest and due items with the same Thema-ID are replaced.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import unicodedata
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document


LANGUAGES = {
    "DEU": {
        "site": "de",
        "card": "Karte",
        "short": "Kurzbeschreibung",
        "article": "Artikel",
        "fields": {
            "Schlagwort": "keyword",
            "Titel": "title",
            "Bibelstelle": "bibleReference",
            "Bibeltext": "bibleText",
        },
    },
    "ENG": {
        "site": "en",
        "card": "Card",
        "short": "Short description",
        "article": "Article",
        "fields": {
            "Keyword": "keyword",
            "Title": "title",
            "Bible reference": "bibleReference",
            "Bible text": "bibleText",
        },
    },
    "FR": {
        "site": "fr",
        "card": "Carte",
        "short": "Description courte",
        "article": "Article",
        "fields": {
            "Mot-clé": "keyword",
            "Titre": "title",
            "Référence biblique": "bibleReference",
            "Texte biblique": "bibleText",
        },
    },
}

DATE_PREFIX = re.compile(r"^(?P<date>\d{6})\s+(?P<topic>.+)\.docx$", re.IGNORECASE)
META_PATTERN = re.compile(
    r"Thema-ID:\s*(?P<theme>.*?)\s*·\s*"
    r"Freigabedatum:\s*(?P<date>\d{2}\.\d{2}\.\d{4})\s*·\s*"
    r"Autor:\s*(?P<author>.+)$"
)
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp"}


class PublishError(RuntimeError):
    """A validation error that should stop publication."""


@dataclass(frozen=True)
class SourceDocument:
    path: Path
    prefix: str
    release_date: date
    filename_topic: str


def slugify(value: str) -> str:
    value = value.replace("ß", "ss").replace("ẞ", "SS")
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^a-z0-9]+", "-", ascii_value.lower()).strip("-")
    if not slug:
        raise PublishError(f"Aus „{value}“ konnte kein URL-Name gebildet werden.")
    return slug


def parse_filename(path: Path) -> SourceDocument | None:
    match = DATE_PREFIX.match(path.name)
    if not match:
        return None
    prefix = match.group("date")
    try:
        release_date = datetime.strptime(prefix, "%y%m%d").date()
    except ValueError as exc:
        raise PublishError(f"Ungültiges Datum im Dateinamen: {path.name}") from exc
    return SourceDocument(
        path=path,
        prefix=prefix,
        release_date=release_date,
        filename_topic=match.group("topic").strip(),
    )


def nonempty_paragraphs(document: Document) -> list[dict[str, str]]:
    return [
        {"text": paragraph.text.strip(), "style": paragraph.style.name}
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]


def parse_metadata(paragraphs: list[dict[str, str]], source: SourceDocument) -> dict[str, str]:
    for paragraph in paragraphs:
        match = META_PATTERN.search(paragraph["text"])
        if match:
            metadata_date = datetime.strptime(match.group("date"), "%d.%m.%Y").date()
            if metadata_date != source.release_date:
                print(
                    "::warning::Das Freigabedatum in der Word-Datei "
                    f"({metadata_date.isoformat()}) weicht vom Dateinamen "
                    f"({source.release_date.isoformat()}) ab. Verwendet wird der Dateiname."
                )
            return {
                "themeId": match.group("theme").strip(),
                "author": match.group("author").strip(),
            }
    raise PublishError(f"Metadatenzeile fehlt in {source.path.name}.")


def find_heading(paragraphs: list[dict[str, str]], text: str, start: int, end: int) -> int:
    for index in range(start, end):
        if paragraphs[index]["text"] == text:
            return index
    raise PublishError(f'Abschnitt „{text}“ fehlt.')


def find_optional_heading(
    paragraphs: list[dict[str, str]], text: str, start: int, end: int
) -> int | None:
    for index in range(start, end):
        if paragraphs[index]["text"] == text:
            return index
    return None


def parse_card_fields(
    paragraphs: list[dict[str, str]],
    start: int,
    end: int,
    field_names: dict[str, str],
) -> dict[str, str]:
    parsed: dict[str, str] = {}
    for paragraph in paragraphs[start:end]:
        text = paragraph["text"]
        if ":" not in text:
            continue
        label, value = text.split(":", 1)
        target = field_names.get(label.strip())
        if target:
            parsed[target] = value.strip()
    missing = [value for value in field_names.values() if not parsed.get(value)]
    if missing:
        raise PublishError(f"Fehlende Kartenfelder: {', '.join(missing)}")
    return parsed


def parse_article_blocks(paragraphs: list[dict[str, str]]) -> list[dict[str, str]]:
    blocks: list[dict[str, str]] = []
    for paragraph in paragraphs:
        text = paragraph["text"]
        style = paragraph["style"]
        if style.startswith("Heading"):
            blocks.append({"type": "heading", "text": text})
        elif "Bible Quote" in style:
            quote, *reference = [part.strip() for part in text.splitlines() if part.strip()]
            block = {"type": "quote", "text": quote}
            if reference:
                block["reference"] = " ".join(reference)
            blocks.append(block)
        else:
            blocks.append({"type": "paragraph", "text": text})
    return blocks


def parse_language(
    paragraphs: list[dict[str, str]],
    code: str,
    start: int,
    end: int,
) -> dict[str, Any]:
    config = LANGUAGES[code]
    card_heading = find_heading(paragraphs, config["card"], start, end)
    short_heading = find_heading(paragraphs, config["short"], card_heading + 1, end)
    article_heading = find_optional_heading(
        paragraphs, config["article"], short_heading + 1, end
    )
    short_end = article_heading if article_heading is not None else end

    card = parse_card_fields(
        paragraphs,
        card_heading + 1,
        short_heading,
        config["fields"],
    )
    short_text = "\n\n".join(
        paragraph["text"] for paragraph in paragraphs[short_heading + 1 : short_end]
    ).strip()
    if not short_text:
        raise PublishError(f"Kurzbeschreibung für {code} fehlt.")

    article_blocks = (
        parse_article_blocks(paragraphs[article_heading + 1 : end])
        if article_heading is not None
        else []
    )
    return {
        **card,
        "shortDescription": short_text,
        "articleBlocks": article_blocks,
        "slug": slugify(card["keyword"]),
    }


def parse_document(source: SourceDocument) -> dict[str, Any]:
    paragraphs = nonempty_paragraphs(Document(source.path))
    metadata = parse_metadata(paragraphs, source)

    language_positions: dict[str, int] = {}
    for index, paragraph in enumerate(paragraphs):
        if paragraph["text"] in LANGUAGES and paragraph["style"].startswith("Heading"):
            language_positions[paragraph["text"]] = index
    missing_languages = sorted(set(LANGUAGES) - set(language_positions))
    if missing_languages:
        raise PublishError(
            f"Sprachabschnitte fehlen in {source.path.name}: {', '.join(missing_languages)}"
        )

    ordered = sorted(language_positions.items(), key=lambda item: item[1])
    languages: dict[str, Any] = {}
    for position, (code, start) in enumerate(ordered):
        end = ordered[position + 1][1] if position + 1 < len(ordered) else len(paragraphs)
        site_language = LANGUAGES[code]["site"]
        languages[site_language] = parse_language(paragraphs, code, start + 1, end)

    return {
        "id": slugify(metadata["themeId"]),
        "themeId": metadata["themeId"],
        "releaseDate": source.release_date.isoformat(),
        "releasePrefix": source.prefix,
        "author": metadata["author"],
        "sourceDocument": source.path.name,
        "languages": languages,
    }


def find_image(input_dir: Path, source: SourceDocument, code: str) -> Path:
    candidates = []
    for path in input_dir.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in IMAGE_SUFFIXES:
            continue
        name_upper = path.stem.upper()
        if path.name.startswith(source.prefix) and re.search(rf"(?:^|\s){code}(?:$|\s)", name_upper):
            candidates.append(path)
    if len(candidates) != 1:
        names = ", ".join(path.name for path in candidates) or "keine"
        raise PublishError(
            f"Für {source.path.name} wurde für {code} nicht genau ein Bild gefunden "
            f"(gefunden: {names})."
        )
    return candidates[0]


def copy_images(
    item: dict[str, Any],
    source: SourceDocument,
    input_dir: Path,
    output_dir: Path,
) -> None:
    item_dir = output_dir / "statusbilder" / item["id"]
    item_dir.mkdir(parents=True, exist_ok=True)
    for code, config in LANGUAGES.items():
        image = find_image(input_dir, source, code)
        target_name = image.name
        target = item_dir / target_name
        shutil.copy2(image, target)
        relative = target.relative_to(output_dir.parent).as_posix()
        item["languages"][config["site"]]["imagePath"] = relative
        item["languages"][config["site"]]["downloadName"] = target_name


def load_manifest(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"schemaVersion": 1, "generatedAt": None, "items": []}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise PublishError(f"Vorhandenes Manifest kann nicht gelesen werden: {path}") from exc
    if data.get("schemaVersion") != 1 or not isinstance(data.get("items"), list):
        raise PublishError(f"Unbekanntes Manifestformat: {path}")
    return data


def discover_due_documents(input_dir: Path, cutoff: date) -> dict[str, tuple[SourceDocument, dict[str, Any]]]:
    latest: dict[str, tuple[SourceDocument, dict[str, Any]]] = {}
    documents = sorted(input_dir.rglob("*.docx"))
    for path in documents:
        source = parse_filename(path)
        if source is None:
            print(f"::notice::Ignoriere Word-Datei ohne Datumspräfix: {path.name}")
            continue
        if source.release_date > cutoff:
            print(f"Später vorgesehen, noch nicht veröffentlicht: {path.name}")
            continue
        item = parse_document(source)
        current = latest.get(item["id"])
        if current is None or current[0].release_date < source.release_date:
            latest[item["id"]] = (source, item)
    return latest


def publish(input_dir: Path, output_dir: Path, manifest_path: Path, cutoff: date) -> int:
    if not input_dir.is_dir():
        raise PublishError(f"Eingangsordner fehlt: {input_dir}")
    output_dir.mkdir(parents=True, exist_ok=True)

    manifest = load_manifest(manifest_path)
    items_by_id = {item["id"]: item for item in manifest["items"]}
    due = discover_due_documents(input_dir, cutoff)

    for item_id, (source, item) in due.items():
        copy_images(item, source, input_dir, output_dir)
        items_by_id[item_id] = item
        print(f"Veröffentlicht/aktualisiert: {item['themeId']} ({item['releaseDate']})")

    items = sorted(
        items_by_id.values(),
        key=lambda item: (item["releaseDate"], item["id"]),
        reverse=True,
    )
    manifest = {
        "schemaVersion": 1,
        "generatedAt": datetime.now(tz=ZoneInfo("Europe/Berlin")).isoformat(),
        "items": items,
    }
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return len(due)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--manifest", type=Path, required=True)
    parser.add_argument(
        "--cutoff",
        help="YYYY-MM-DD; standardmäßig das heutige Datum in Europe/Berlin",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    cutoff = (
        datetime.strptime(args.cutoff, "%Y-%m-%d").date()
        if args.cutoff
        else datetime.now(tz=ZoneInfo("Europe/Berlin")).date()
    )
    try:
        count = publish(args.input, args.output, args.manifest, cutoff)
    except PublishError as exc:
        print(f"::error::{exc}", file=sys.stderr)
        return 1
    print(f"Fertig. {count} fällige Themen verarbeitet.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
